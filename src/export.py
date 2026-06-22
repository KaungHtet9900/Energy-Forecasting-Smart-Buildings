"""Export paper/manuscript.md to DOCX (python-docx) and PDF (reportlab).

Handles headings, paragraphs with **bold**/*italic*, pipe tables, embedded figures
(resolved relative to paper/), and italic figure captions. No system dependencies.

Usage:  python -m src.export
"""
from __future__ import annotations

import re
from pathlib import Path

from src.utils import load_config, get_logger

log = get_logger("export", logfile="logs/export.log")

IMG_RE = re.compile(r"!\[(.*?)\]\((.*?)\)")
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


# --------------------------------------------------------------------------
# Markdown -> blocks
# --------------------------------------------------------------------------
def parse_blocks(md_text: str):
    lines = md_text.split("\n")
    blocks, i = [], 0
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            i += 1
            continue
        if s.startswith("#"):
            level = len(s) - len(s.lstrip("#"))
            blocks.append(("h", level, s.lstrip("#").strip()))
            i += 1
        elif s.startswith("!["):
            m = IMG_RE.match(s)
            if m:
                blocks.append(("img", m.group(2), m.group(1)))
            i += 1
        elif s.startswith("|"):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i].strip())
                i += 1
            blocks.append(("table", tbl))
        else:
            blocks.append(("p", s))
            i += 1
    return blocks


def parse_table(tbl_lines):
    rows = [[c.strip() for c in ln.strip().strip("|").split("|")] for ln in tbl_lines]
    header = rows[0]
    data = [r for r in rows[1:] if not all(set(c) <= set("-: ") for c in r)]
    return header, data


def resolve_img(base_dir: Path, rel: str) -> Path:
    return (base_dir / rel).resolve()


# --------------------------------------------------------------------------
# DOCX
# --------------------------------------------------------------------------
def build_docx(cfg, blocks, base_dir: Path, out_path: Path):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    def rich_paragraph(text, style=None, italic=False, size=None, center=False):
        p = doc.add_paragraph(style=style)
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for part in re.split(r"(\*\*.+?\*\*)", text):
            if part.startswith("**") and part.endswith("**"):
                r = p.add_run(part[2:-2]); r.bold = True
            else:
                r = p.add_run(part)
            if italic:
                r.italic = True
            if size:
                r.font.size = Pt(size)
        return p

    for blk in blocks:
        kind = blk[0]
        if kind == "h":
            level = blk[1]
            doc.add_heading(blk[2], level=max(0, level - 1))
        elif kind == "p":
            text = blk[1]
            if text.startswith("*") and text.endswith("*") and not text.startswith("**"):
                rich_paragraph(text.strip("*"), italic=True, size=9, center=True)
            else:
                rich_paragraph(text)
        elif kind == "img":
            path = resolve_img(base_dir, blk[1])
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                log.warning("missing image: %s", path)
        elif kind == "table":
            header, data = parse_table(blk[1])
            t = doc.add_table(rows=1, cols=len(header))
            try:
                t.style = "Light Grid Accent 1"
            except Exception:  # noqa: BLE001
                t.style = "Table Grid"
            for j, h in enumerate(header):
                run = t.rows[0].cells[j].paragraphs[0].add_run(h)
                run.bold = True
            for row in data:
                cells = t.add_row().cells
                for j, c in enumerate(row[:len(header)]):
                    cells[j].text = c
    doc.save(str(out_path))
    log.info("wrote %s", out_path)


# --------------------------------------------------------------------------
# PDF
# --------------------------------------------------------------------------
def _inline_html(text: str) -> str:
    text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    text = BOLD_RE.sub(r"<b>\1</b>", text)
    text = re.sub(r"(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)", r"<i>\1</i>", text)
    return text


def build_pdf(cfg, blocks, base_dir: Path, out_path: Path):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Image,
                                    Table, TableStyle)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from PIL import Image as PILImage

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle("Cap", parent=styles["Normal"], fontSize=8,
                              textColor=colors.grey, alignment=1, spaceAfter=8))
    body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=9.5, leading=13,
                          spaceAfter=6, alignment=4)
    h_styles = {1: styles["Title"],
                2: ParagraphStyle("H2", parent=styles["Heading1"], fontSize=14, spaceBefore=10),
                3: ParagraphStyle("H3", parent=styles["Heading2"], fontSize=11.5, spaceBefore=6)}
    max_w = 6.5 * inch

    story = []
    for blk in blocks:
        kind = blk[0]
        if kind == "h":
            story.append(Paragraph(_inline_html(blk[2]), h_styles.get(blk[1], h_styles[3])))
        elif kind == "p":
            text = blk[1]
            if text.startswith("*") and text.endswith("*") and not text.startswith("**"):
                story.append(Paragraph(_inline_html(text.strip("*")), styles["Cap"]))
            else:
                story.append(Paragraph(_inline_html(text), body))
        elif kind == "img":
            path = resolve_img(base_dir, blk[1])
            if path.exists():
                iw, ih = PILImage.open(path).size
                w = max_w
                h = w * ih / iw
                if h > 4.3 * inch:
                    h = 4.3 * inch
                    w = h * iw / ih
                story.append(Spacer(1, 6))
                story.append(Image(str(path), width=w, height=h, hAlign="CENTER"))
                story.append(Spacer(1, 2))
            else:
                log.warning("missing image: %s", path)
        elif kind == "table":
            header, data = parse_table(blk[1])
            cell = ParagraphStyle("Cell", parent=styles["Normal"], fontSize=7.5, leading=9)
            head = ParagraphStyle("Head", parent=cell, textColor=colors.white,
                                  fontName="Helvetica-Bold")
            tdata = [[Paragraph(_inline_html(h), head) for h in header]]
            tdata += [[Paragraph(_inline_html(c), cell) for c in row[:len(header)]]
                      for row in data]
            t = Table(tdata, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2F5597")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EEF2FA")]),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#B0B0B0")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 2),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
            ]))
            story.append(t)
            story.append(Spacer(1, 8))

    doc = SimpleDocTemplate(str(out_path), pagesize=letter,
                            leftMargin=0.9 * inch, rightMargin=0.9 * inch,
                            topMargin=0.9 * inch, bottomMargin=0.9 * inch,
                            title="Interpretable Attention-Based Building Energy Forecasting")
    doc.build(story)
    log.info("wrote %s", out_path)


def main() -> int:
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("--config", default=None)
    args = ap.parse_args()
    cfg = load_config(args.config)

    paper_dir = cfg["project_root"] / "paper"
    md_path = paper_dir / "manuscript.md"
    md_text = md_path.read_text(encoding="utf-8")
    blocks = parse_blocks(md_text)
    log.info("parsed %d blocks from %s", len(blocks), md_path)

    build_docx(cfg, blocks, paper_dir, paper_dir / "manuscript.docx")
    build_pdf(cfg, blocks, paper_dir, paper_dir / "manuscript.pdf")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
